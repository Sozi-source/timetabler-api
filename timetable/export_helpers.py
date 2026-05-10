"""
timetable/export_helpers.py
============================
PDF and DOCX timetable builders.

Dependencies (add to requirements.txt):
    reportlab>=4.0
    python-docx>=1.1
"""

from __future__ import annotations

import io
from typing import Any


def _cell_lines(item: dict | None, is_master: bool) -> list[str]:
    if item is None:
        return []
    if isinstance(item, list):
        lines: list[str] = []
        for e in item:
            lines.append(f"{e.get('unit_code', '')} – {e.get('unit_name', '')}")
            if e.get("trainer_name"):
                lines.append(f"  {e['trainer_name']}")
            parts = [e.get("room_code", ""), e.get("cohort_name", "")]
            lines.append("  " + "  ·  ".join(p for p in parts if p))
            lines.append("")
        return [l for l in lines if l.strip() or lines.index(l) < len(lines) - 1]
    e = item
    lines = [f"{e.get('unit_code', '')} – {e.get('unit_name', '')}"]
    if e.get("trainer_name"):
        lines.append(e["trainer_name"])
    if e.get("room_code"):
        lines.append(e["room_code"])
    return lines


# ─────────────────────────────────────────────────────────────────────────────
# PDF builder  (reportlab)
# ─────────────────────────────────────────────────────────────────────────────

def build_pdf_timetable(
    title: str,
    subtitle: str,
    days: list[str],
    periods: list[dict],
    grid: dict[str, dict[str, Any]],
    is_master: bool = False,
) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A3, landscape
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    HEADER_BG   = colors.HexColor("#16213e")
    HEADER_FG   = colors.white
    DAY_BG      = colors.HexColor("#2e75b6")
    GROUP_BG    = colors.HexColor("#d6e4f0")
    CELL_BG     = colors.HexColor("#ffffff")
    ALT_BG      = colors.HexColor("#f0f4ff")
    BORDER      = colors.HexColor("#aaaaaa")
    TITLE_COLOR = colors.HexColor("#1f3864")
    SUB_COLOR   = colors.HexColor("#5a6180")

    title_style = ParagraphStyle("TT", fontName="Helvetica-Bold", fontSize=14,
                                 textColor=TITLE_COLOR, spaceAfter=2)
    sub_style   = ParagraphStyle("TS", fontName="Helvetica", fontSize=9,
                                 textColor=SUB_COLOR, spaceAfter=6)
    header_style = ParagraphStyle("TH", fontName="Helvetica-Bold", fontSize=7,
                                  textColor=HEADER_FG, leading=10, alignment=1)
    day_style    = ParagraphStyle("TD", fontName="Helvetica-Bold", fontSize=8,
                                  textColor=colors.white, alignment=1)
    group_style  = ParagraphStyle("TG", fontName="Helvetica-Bold", fontSize=7,
                                  textColor=colors.HexColor("#1f3864"), alignment=1)
    code_style   = ParagraphStyle("TC", fontName="Helvetica-Bold", fontSize=7,
                                  textColor=colors.HexColor("#1f3864"), leading=9, spaceAfter=1)
    detail_style = ParagraphStyle("TDet", fontName="Helvetica", fontSize=6,
                                  textColor=SUB_COLOR, leading=8)

    buf  = io.BytesIO()
    page = landscape(A3)
    doc  = SimpleDocTemplate(buf, pagesize=page,
                             leftMargin=12*mm, rightMargin=12*mm,
                             topMargin=12*mm, bottomMargin=12*mm)

    def _slot_content(entries):
        if not entries:
            return [Paragraph("", detail_style)]
        paras = []
        items = entries if isinstance(entries, list) else [entries]
        for e in items:
            paras.append(Paragraph(
                f"<b>{e.get('unit_code','')}</b> {e.get('unit_name','')}",
                code_style))
            info_parts = [e.get("trainer_name",""), e.get("room_code","")]
            if is_master and e.get("cohort_name"):
                info_parts.append(e["cohort_name"])
            info = "  ·  ".join(p for p in info_parts if p)
            if info:
                paras.append(Paragraph(info, detail_style))
            paras.append(Spacer(1, 2))
        return paras

    # Collect all cohort/group names in order from the grid
    # For master: each cell is a list of entries — cohort names are within entries
    # We need to build rows: DAY | GROUP | slot1 | slot2 | ... per group per day
    # Gather unique groups across all days
    group_order: list[str] = []
    seen: set[str] = set()
    for day in days:
        for pid in [p["id"] for p in periods]:
            item = grid.get(day, {}).get(pid)
            if not item:
                continue
            entries = item if isinstance(item, list) else [item]
            for e in entries:
                g = e.get("cohort_name", "") or e.get("trainer_name", "") or "—"
                if g not in seen:
                    seen.add(g)
                    group_order.append(g)

    # Header row: DAYS | INTAKE GROUP | period labels...
    header_row = [
        Paragraph("DAYS", header_style),
        Paragraph("INTAKE GROUP", header_style),
    ] + [
        Paragraph(f"<b>{p['label']}</b><br/>{p['start'][:5]}–{p['end'][:5]}", header_style)
        for p in periods
    ]

    table_data = [header_row]

    # One row per (day, group)
    row_index = 0
    span_info = []  # for SPAN commands later
    for day in days:
        groups_in_day = []
        # collect groups that appear in this day
        day_groups_seen: list[str] = []
        day_groups_set: set[str] = set()
        for pid in [p["id"] for p in periods]:
            item = grid.get(day, {}).get(pid)
            if not item:
                continue
            entries = item if isinstance(item, list) else [item]
            for e in entries:
                g = e.get("cohort_name", "") or e.get("trainer_name", "") or "—"
                if g not in day_groups_set:
                    day_groups_set.add(g)
                    day_groups_seen.append(g)

        if not day_groups_seen:
            day_groups_seen = ["—"]

        n = len(day_groups_seen)
        span_info.append((row_index + 1, n))  # +1 for header

        for gi, group in enumerate(day_groups_seen):
            row = []
            if gi == 0:
                row.append(Paragraph(day, day_style))
            else:
                row.append(Paragraph("", day_style))  # placeholder — merged via SPAN

            row.append(Paragraph(group, group_style))

            for p in periods:
                item = grid.get(day, {}).get(p["id"])
                if not item:
                    row.append(Paragraph("", detail_style))
                    continue
                entries = item if isinstance(item, list) else [item]
                # filter entries for this group
                group_entries = [
                    e for e in entries
                    if (e.get("cohort_name") or e.get("trainer_name") or "—") == group
                ]
                row.append(_slot_content(group_entries) if group_entries else [Paragraph("", detail_style)])

            table_data.append(row)
            row_index += 1

    n_cols = 2 + len(periods)
    usable_width = page[0] - 24 * mm
    day_w   = 18 * mm
    group_w = 28 * mm
    slot_w  = (usable_width - day_w - group_w) / max(len(periods), 1)
    col_widths = [day_w, group_w] + [slot_w] * len(periods)

    tbl = Table(table_data, colWidths=col_widths, repeatRows=1)

    style_cmds = [
        ("BACKGROUND",    (0, 0), (-1, 0),  HEADER_BG),
        ("TEXTCOLOR",     (0, 0), (-1, 0),  HEADER_FG),
        ("ALIGN",         (0, 0), (-1, 0),  "CENTER"),
        ("VALIGN",        (0, 0), (-1, 0),  "MIDDLE"),
        ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",      (0, 0), (-1, 0),  7),
        ("BOTTOMPADDING", (0, 0), (-1, 0),  5),
        ("TOPPADDING",    (0, 0), (-1, 0),  5),
        ("BACKGROUND",    (1, 1), (1, -1),  GROUP_BG),
        ("ALIGN",         (1, 1), (1, -1),  "CENTER"),
        ("VALIGN",        (0, 1), (-1, -1), "TOP"),
        ("BACKGROUND",    (2, 1), (-1, -1), CELL_BG),
        ("GRID",          (0, 0), (-1, -1), 0.5, BORDER),
        ("LEFTPADDING",   (0, 0), (-1, -1), 4),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 4),
        ("TOPPADDING",    (1, 1), (-1, -1), 3),
        ("BOTTOMPADDING", (1, 1), (-1, -1), 3),
    ]

    # Day column spans and colour
    for start_row, count in span_info:
        end_row = start_row + count - 1
        style_cmds.append(("BACKGROUND", (0, start_row), (0, end_row), DAY_BG))
        style_cmds.append(("ALIGN",      (0, start_row), (0, end_row), "CENTER"))
        style_cmds.append(("VALIGN",     (0, start_row), (0, end_row), "MIDDLE"))
        style_cmds.append(("FONTNAME",   (0, start_row), (0, end_row), "Helvetica-Bold"))
        style_cmds.append(("TEXTCOLOR",  (0, start_row), (0, end_row), colors.white))
        if count > 1:
            style_cmds.append(("SPAN", (0, start_row), (0, end_row)))

    tbl.setStyle(TableStyle(style_cmds))

    story = [
        Paragraph(title, title_style),
        Paragraph(subtitle, sub_style),
        tbl,
    ]
    doc.build(story)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# DOCX builder  (python-docx)  — matches uploaded HND timetable format
# ─────────────────────────────────────────────────────────────────────────────

def build_docx_timetable(
    title: str,
    subtitle: str,
    days: list[str],
    periods: list[dict],
    grid: dict[str, dict[str, Any]],
    is_master: bool = False,
) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor, Twips
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

    # ── Colours ───────────────────────────────────────────────────────────────
    HEADER_BG = "1F3864"
    DAY_BG    = "2E75B6"
    GROUP_BG  = "D6E4F0"
    CELL_BG   = "FFFFFF"
    ALT_BG    = "F0F4FF"
    WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
    DARK      = RGBColor(0x1F, 0x38, 0x64)
    GREY      = RGBColor(0x44, 0x44, 0x44)
    BORDER_C  = "AAAAAA"

    def _set_bg(cell, hex_color: str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = parse_xml(
            f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
        )
        tcPr.append(shd)

    def _set_borders(cell, color: str = "AAAAAA", sz: int = 4):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        xml  = (
            f'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:top    w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:left   w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'<w:right  w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(parse_xml(xml))

    def _run(para, text, bold=False, size_pt=8, color=None, italic=False):
        run = para.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size_pt)
        if color:
            run.font.color.rgb = color
        return run

    def _fill_slot(cell, entries, is_master):
        """Fill a time-slot cell with one or more scheduled unit entries."""
        first_para = cell.paragraphs[0]
        first_para.paragraph_format.space_after  = Pt(0)
        first_para.paragraph_format.space_before = Pt(0)

        if not entries:
            return

        items = entries if isinstance(entries, list) else [entries]
        first = True
        for e in items:
            if first:
                p = first_para
                first = False
            else:
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after  = Pt(0)

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            unit = f"{e.get('unit_name','') or e.get('unit_code','')}"
            _run(p, unit, bold=True, size_pt=8, color=DARK)

            details = []
            if e.get("trainer_name"): details.append(e["trainer_name"])
            if e.get("room_code"):    details.append(e["room_code"])
            if is_master and e.get("cohort_name"): details.append(e["cohort_name"])

            if details:
                dp = cell.add_paragraph()
                dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                dp.paragraph_format.space_before = Pt(0)
                dp.paragraph_format.space_after  = Pt(0)
                _run(dp, "  ·  ".join(details), size_pt=7, color=GREY)

    # ── Document ──────────────────────────────────────────────────────────────
    doc     = Document()
    section = doc.sections[0]
    section.page_width   = Cm(42.0)
    section.page_height  = Cm(29.7)
    section.left_margin  = section.right_margin  = Cm(1.0)
    section.top_margin   = section.bottom_margin = Cm(1.0)
    section.orientation  = 1  # LANDSCAPE

    # Title
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(2)
    r = tp.add_run(title)
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = DARK

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(6)
    r = sp.add_run(subtitle)
    r.font.size = Pt(9); r.font.color.rgb = GREY

    # ── Collect groups per day ────────────────────────────────────────────────
    # groups_per_day[day] = ordered list of unique group names
    groups_per_day: dict[str, list[str]] = {}
    for day in days:
        seen_g: set[str] = set()
        ordered_g: list[str] = []
        for p in periods:
            item = grid.get(day, {}).get(p["id"])
            if not item:
                continue
            entries = item if isinstance(item, list) else [item]
            for e in entries:
                g = e.get("cohort_name") or e.get("trainer_name") or "—"
                if g not in seen_g:
                    seen_g.add(g)
                    ordered_g.append(g)
        groups_per_day[day] = ordered_g if ordered_g else ["—"]

    # ── Table dimensions ──────────────────────────────────────────────────────
    # Columns: DAY | INTAKE GROUP | period1 | period2 | ...
    n_cols     = 2 + len(periods)
    total_rows = 1 + sum(len(groups_per_day[d]) for d in days)  # header + data

    tbl = doc.add_table(rows=total_rows, cols=n_cols)
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Column widths (DXA: 1440 = 1 inch, Cm(1) = 567 twips)
    usable_cm = 42.0 - 2.0          # page minus margins
    day_cm    = 2.2
    group_cm  = 3.0
    slot_cm   = (usable_cm - day_cm - group_cm) / max(len(periods), 1)

    def _set_col_width(cell, cm):
        cell.width = Cm(cm)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = parse_xml(
            f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            f' w:w="{int(cm * 567)}" w:type="dxa"/>'
        )
        # remove existing tcW if present
        existing = tcPr.find(qn('w:tcW'))
        if existing is not None:
            tcPr.remove(existing)
        tcPr.append(tcW)

    # ── Header row ────────────────────────────────────────────────────────────
    hdr_row = tbl.rows[0]
    hdr_cells = hdr_row.cells

    for ci, (cell, label, cm) in enumerate(zip(
        hdr_cells,
        ["DAYS", "INTAKE GROUP"] + [f"{p['label']}\n{p['start'][:5]}–{p['end'][:5]}" for p in periods],
        [day_cm, group_cm] + [slot_cm] * len(periods),
    )):
        _set_col_width(cell, cm)
        _set_bg(cell, HEADER_BG)
        _set_borders(cell, BORDER_C)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        lines = label.split("\n")
        _run(p, lines[0], bold=True, size_pt=8, color=WHITE)
        if len(lines) > 1:
            p.add_run("\n")
            _run(p, lines[1], bold=False, size_pt=7, color=RGBColor(0xCC, 0xD0, 0xFF))

    # ── Data rows ─────────────────────────────────────────────────────────────
    ri = 1
    for day in days:
        groups = groups_per_day[day]
        n_g    = len(groups)

        for gi, group in enumerate(groups):
            row = tbl.rows[ri]

            # ── DAY cell (only first row of this day, merge the rest) ─────────
            day_cell = row.cells[0]
            _set_col_width(day_cell, day_cm)
            _set_bg(day_cell, DAY_BG)
            _set_borders(day_cell, BORDER_C)
            day_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            if gi == 0:
                p = day_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(0)
                _run(p, day, bold=True, size_pt=9, color=WHITE)

            # ── GROUP cell ───────────────────────────────────────────────────
            grp_cell = row.cells[1]
            _set_col_width(grp_cell, group_cm)
            _set_bg(grp_cell, GROUP_BG)
            _set_borders(grp_cell, BORDER_C)
            grp_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = grp_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            _run(p, group, bold=True, size_pt=8, color=DARK)

            # ── Slot cells ───────────────────────────────────────────────────
            for ci, period in enumerate(periods, start=2):
                sc = row.cells[ci]
                _set_col_width(sc, slot_cm)
                _set_bg(sc, ALT_BG if ri % 2 == 0 else CELL_BG)
                _set_borders(sc, BORDER_C)
                sc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                item = grid.get(day, {}).get(period["id"])
                entries = []
                if item:
                    all_entries = item if isinstance(item, list) else [item]
                    entries = [
                        e for e in all_entries
                        if (e.get("cohort_name") or e.get("trainer_name") or "—") == group
                    ]
                _fill_slot(sc, entries, is_master)

            ri += 1

        # ── Merge the DAY column cells for this day ───────────────────────────
        if n_g > 1:
            first_day_cell = tbl.rows[ri - n_g].cells[0]
            last_day_cell  = tbl.rows[ri - 1].cells[0]
            first_day_cell.merge(last_day_cell)
            # Re-apply styles after merge
            _set_bg(first_day_cell, DAY_BG)
            _set_borders(first_day_cell, BORDER_C)
            first_day_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ── Serialise ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()